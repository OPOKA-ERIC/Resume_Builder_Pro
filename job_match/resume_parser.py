import os
import re
import subprocess
import tempfile
from pathlib import Path

from django.core.files.uploadedfile import UploadedFile


def extract_text_from_upload(file_obj: UploadedFile) -> str:
    ext = Path(file_obj.name).suffix.lower() if file_obj.name else '.txt'

    if ext == '.pdf':
        return _extract_pdf(file_obj)
    elif ext == '.docx':
        return _extract_docx(file_obj)
    elif ext in ('.txt', '.text', '.md', ''):
        return _extract_text(file_obj)
    else:
        return _extract_text(file_obj)


def _extract_pdf(file_obj: UploadedFile) -> str:
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        for chunk in file_obj.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ['pdftotext', '-layout', tmp_path, '-'],
            capture_output=True, text=True, timeout=30
        )
        text = result.stdout.strip()
        if not text:
            result = subprocess.run(
                ['pdftotext', tmp_path, '-'],
                capture_output=True, text=True, timeout=30
            )
            text = result.stdout.strip()
        return text
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
        return _fallback_pdf(tmp_path)
    finally:
        os.unlink(tmp_path)


def _fallback_pdf(path: str) -> str:
    try:
        with open(path, 'rb') as f:
            data = f.read()
        text_parts = []
        raw = data.decode('latin-1')
        for match in re.finditer(r'\((.*?)\)', raw):
            text_parts.append(match.group(1))
        for match in re.finditer(r'<([0-9A-Fa-f]+)>', raw):
            try:
                text_parts.append(bytes.fromhex(match.group(1)).decode('latin-1'))
            except Exception:
                pass
        result = ' '.join(text_parts)
        result = re.sub(r'\\[0-9]{3}', ' ', result)
        return re.sub(r'\s+', ' ', result).strip()
    except Exception:
        return ''


def _extract_docx(file_obj: UploadedFile) -> str:
    import docx
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        for chunk in file_obj.chunks():
            tmp.write(chunk)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return '\n'.join(p.strip() for p in paragraphs if p.strip())
    except Exception:
        return ''
    finally:
        os.unlink(tmp_path)


def _extract_text(file_obj: UploadedFile) -> str:
    try:
        return file_obj.read().decode('utf-8', errors='replace')
    except Exception:
        return ''
